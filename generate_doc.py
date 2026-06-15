"""Generates a Word (.docx) project documentation file with screenshots."""
import os
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def code_block(text):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)

def table_row(tbl, cells):
    row = tbl.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = val
    return row

def add_screenshot(path, caption="", width=Inches(5.5)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        return True
    return False

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Face Recognition Attendance System")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("Project Documentation")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = info.add_run(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}")
run3.font.size = Pt(11)

doc.add_paragraph()
auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = auth.add_run("Developed using Python  •  Streamlit  •  face_recognition  •  SQLite")
run4.font.size = Pt(10)
run4.italic = True
run4.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
heading("Table of Contents", 1)
toc_items = [
    ("1.", "Project Overview"),
    ("2.", "Tech Stack"),
    ("3.", "Project Structure"),
    ("4.", "Database Schema"),
    ("5.", "Module Descriptions"),
    ("6.", "Page-by-Page Breakdown"),
    ("7.", "Key Code Snippets"),
    ("8.", "Location Restriction Feature"),
    ("9.", "Application Screenshots"),
]
for num, item in toc_items:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{num}  {item}").font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Project Overview")
para(
    "The Face Recognition Attendance System is a Python-based web application that automates "
    "student attendance using real-time facial recognition. Built with Streamlit, it provides "
    "a complete end-to-end solution for universities — from student enrollment and live face "
    "detection to check-in/check-out recording, GPS-based location verification, and detailed "
    "report generation."
)
doc.add_paragraph()

heading("Key Features", 2)
features = [
    "Role-based login — Student, Teacher, and Admin accounts with separate views",
    "Camera-based face enrollment with 3–5 training photos per student",
    "Multi-face detection in a single frame — mark an entire class at once",
    "Check-in and check-out time recording with timestamp",
    "GPS location restriction — blocks attendance if device is outside configured radius",
    "Date-range and monthly attendance reports with interactive Plotly charts",
    "CSV export for all reports",
    "SQLite database — zero external database server required",
    "Profile photo and training dataset stored per student",
    "Attendance percentage tracking with 75% threshold alert",
]
for f in features:
    doc.add_paragraph(f, style="List Bullet")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. TECH STACK
# ══════════════════════════════════════════════════════════════════════════════
heading("2. Tech Stack")

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
hdr[0].text = "Component"
hdr[1].text = "Library / Tool"
hdr[2].text = "Version / Notes"
for cell in hdr:
    for p_ in cell.paragraphs:
        for r in p_.runs:
            r.bold = True

rows = [
    ("UI Framework",        "Streamlit",             ">= 1.35.0"),
    ("Face Recognition",    "face_recognition",      ">= 1.3.0  (dlib backend)"),
    ("Computer Vision",     "OpenCV",                ">= 4.9.0"),
    ("Numerical Computing", "NumPy",                 ">= 1.26.0"),
    ("Image Processing",    "Pillow",                ">= 10.3.0"),
    ("Data Analysis",       "Pandas",                ">= 2.2.0"),
    ("Charting",            "Plotly",                ">= 5.22.0"),
    ("Geolocation",         "streamlit-geolocation", ">= 0.0.10"),
    ("Database",            "SQLite3",               "Built-in Python stdlib"),
    ("Language",            "Python",                "3.11"),
]
for r in rows:
    table_row(tbl, r)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
heading("3. Project Structure")
code_block("""\
face-recognition-attendance/
│
├── streamlit_app.py            ← Main app entry point (login + dashboard)
│
├── pages/
│   ├── 1_Register_Face.py      ← Enroll students + capture training photos
│   ├── 2_Mark_Attendance.py    ← Camera-based check-in / check-out
│   ├── 3_Manage_Users.py       ← View, retrain, delete students
│   ├── 4_Reports.py            ← Date-range & monthly reports + CSV export
│   ├── 5_Data_Viewer.py        ← Raw data browser (students, records, photos)
│   └── 6_Settings.py           ← GPS location restriction configuration
│
├── utils/
│   ├── db_utils.py             ← SQLite helpers + location + Haversine distance
│   └── face_utils.py           ← face_recognition wrappers (encode, detect, draw)
│
├── data/
│   ├── attendance.db           ← SQLite database (auto-created)
│   └── dataset/                ← Training photos saved per student
│       └── <roll>_<name>/
│           ├── photo_1.jpg
│           └── ...
│
├── backend/                    ← Separate Flask REST API skeleton
│   └── app/
│       ├── routes/             ← auth, attendance, reports, locations
│       ├── models/             ← user, attendance, location
│       └── services/           ← face_recognition, attendance_service
│
├── screenshots/                ← UI screenshots
├── requirements_streamlit.txt  ← Python dependencies
└── generate_doc.py             ← This documentation generator
""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. DATABASE SCHEMA
# ══════════════════════════════════════════════════════════════════════════════
heading("4. Database Schema")
para("The application uses a single SQLite file at data/attendance.db with three main tables.")
doc.add_paragraph()

heading("4.1  users", 2)
code_block("""\
CREATE TABLE users (
    id              INTEGER PRIMARY KEY AUTOINCREMENT,
    name            TEXT    NOT NULL,
    employee_id     TEXT    UNIQUE NOT NULL,   -- Roll number / Student ID
    email           TEXT,
    department      TEXT,                      -- Course / Branch
    semester        TEXT,
    year            TEXT,
    password_hash   TEXT,                      -- Hashed login password
    face_encodings  BLOB,                      -- Pickle of [encoding, ...]
    profile_photo   BLOB,                      -- JPEG bytes of profile picture
    created_at      TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    is_active       BOOLEAN DEFAULT 1
);""")

doc.add_paragraph()
heading("4.2  attendance", 2)
code_block("""\
CREATE TABLE attendance (
    id              INTEGER PRIMARY KEY AUTOINCREMENT,
    user_id         INTEGER NOT NULL REFERENCES users(id),
    date            DATE    NOT NULL,
    check_in_time   TIMESTAMP,
    check_out_time  TIMESTAMP,
    location        TEXT    DEFAULT 'Lecture Hall',
    status          TEXT    DEFAULT 'Present'
);""")

doc.add_paragraph()
heading("4.3  location_settings", 2)
code_block("""\
CREATE TABLE location_settings (
    id              INTEGER PRIMARY KEY CHECK (id = 1),  -- singleton row
    latitude        REAL,
    longitude       REAL,
    radius_meters   INTEGER DEFAULT 100,
    location_name   TEXT    DEFAULT 'Campus'
);""")

doc.add_paragraph()
heading("4.4  teachers", 2)
code_block("""\
CREATE TABLE teachers (
    id              INTEGER PRIMARY KEY AUTOINCREMENT,
    name            TEXT    NOT NULL,
    teacher_id      TEXT    UNIQUE NOT NULL,
    email           TEXT,
    department      TEXT,
    password_hash   TEXT,
    is_approved     BOOLEAN DEFAULT 0,         -- Admin must approve before login
    created_at      TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. MODULE DESCRIPTIONS
# ══════════════════════════════════════════════════════════════════════════════
heading("5. Module Descriptions")

heading("5.1  utils/db_utils.py", 2)
tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = "Table Grid"
tbl2.rows[0].cells[0].text = "Function"
tbl2.rows[0].cells[1].text = "Description"
for c in tbl2.rows[0].cells:
    for p_ in c.paragraphs:
        for r in p_.runs:
            r.bold = True

db_fns = [
    ("init_db()", "Creates all tables and migrates missing columns on first run"),
    ("create_user(...)", "Inserts a new student; returns new ID or None on duplicate roll number"),
    ("get_all_users()", "Returns all active students ordered by name"),
    ("get_user_by_id(id)", "Fetch a single user dict by primary key"),
    ("delete_user(id)", "Soft-delete: sets is_active = 0"),
    ("save_face_encodings_db(id, list)", "Serialise and store face encodings as BLOB"),
    ("load_all_face_encodings_db()", "Load {user_id: [enc...]} for all active enrolled students"),
    ("save_profile_photo(id, bytes)", "Persist JPEG bytes in users table"),
    ("get_profile_photo(id)", "Return JPEG bytes or None"),
    ("save_dataset_photos(...)", "Save training photos to data/dataset/<roll>_<name>/"),
    ("mark_checkin(user_id, location)", "Insert attendance row; returns ('success'|'already_in'|'already_done', id)"),
    ("mark_checkout(user_id)", "Set check_out_time on today's open row; returns bool"),
    ("get_today_attendance()", "All attendance rows joined with user info for today"),
    ("get_attendance_by_date_range(...)", "Filter attendance by date range and optional student"),
    ("get_attendance_summary(year, month)", "Per-student days-present count for a month"),
    ("haversine_distance(lat1, lon1, lat2, lon2)", "Return distance in metres between two GPS coordinates"),
    ("create_teacher(...)", "Register a teacher account (pending admin approval)"),
    ("verify_student_login(roll, pwd)", "Validate student credentials; returns user dict or None"),
    ("verify_teacher_login(tid, pwd)", "Validate teacher credentials; returns teacher dict or None"),
]
for fn, desc in db_fns:
    table_row(tbl2, [fn, desc])

doc.add_paragraph()
heading("5.2  utils/face_utils.py", 2)
tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = "Table Grid"
tbl3.rows[0].cells[0].text = "Function"
tbl3.rows[0].cells[1].text = "Description"
for c in tbl3.rows[0].cells:
    for p_ in c.paragraphs:
        for r in p_.runs:
            r.bold = True

face_fns = [
    ("load_encodings()", "Load all face encodings from database"),
    ("encode_image(image)", "Detect exactly one face; return (encoding, None) or (None, error_msg)"),
    ("save_face_encodings_bulk(id, list)", "Replace stored encodings with up to 5 samples"),
    ("recognize_faces(image, tolerance)", "Identify all faces in frame; returns list of {user_id, location, confidence}"),
    ("draw_annotations(image, results, user_map)", "Draw bounding boxes and name labels on image"),
]
for fn, desc in face_fns:
    table_row(tbl3, [fn, desc])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. PAGE-BY-PAGE BREAKDOWN
# ══════════════════════════════════════════════════════════════════════════════
heading("6. Page-by-Page Breakdown")

pages_desc = [
    (
        "streamlit_app.py — Login & Dashboard",
        "The entry point. Shows a login page with tabs for Student and Teacher login, plus a "
        "registration form for new accounts. After login, the Teacher/Admin sees a dashboard "
        "with four metrics (Enrolled Students, Faces Trained, Present Today, Absent Today) and "
        "a live table of today's attendance. Students see their own personal attendance summary."
    ),
    (
        "1_Register_Face.py — Register New Student",
        "Split layout: student details form on the left, live camera on the right. "
        "The teacher captures 3–5 face photos per student. Each photo is encoded and stored. "
        "The first photo becomes the profile picture. All photos are saved to disk under "
        "data/dataset/<roll>_<name>/ for audit and retraining purposes. "
        "A grid of enrolled students with their profile photos is shown at the bottom."
    ),
    (
        "2_Mark_Attendance.py — Mark Attendance",
        "Camera-based check-in / check-out. The teacher takes a photo; the system recognises "
        "all faces simultaneously. Each recognised student is marked Present or given an Exit "
        "record. Students can also self check-in using their own camera. "
        "Includes GPS location verification — if a reference location is configured, "
        "the device must be within the allowed radius."
    ),
    (
        "3_Manage_Users.py — Manage Students",
        "Lists all enrolled students with search/filter. Each card shows profile photo, "
        "student info, training status, and action buttons: Retrain Face (new camera session), "
        "Remove Face Data, and Delete Student (soft-delete sets is_active = 0)."
    ),
    (
        "4_Reports.py — Attendance Reports",
        "Two tabs: (1) Date Range Report — filter by date and student, view a table plus a "
        "daily bar chart, download CSV. (2) Monthly Summary — per-student days-present count "
        "with a colour-coded Plotly bar chart and a 75% attendance threshold reference line. "
        "Students below 75% are highlighted."
    ),
    (
        "5_Data_Viewer.py — Data Viewer",
        "Raw data browser with three tabs: Students (paginated, with toggle for deleted records), "
        "Attendance Records (filterable by student and date, full CSV export), "
        "Dataset Photos (grid view of all training images per student folder)."
    ),
    (
        "6_Settings.py — Settings",
        "Admin configuration page for the GPS location restriction feature. "
        "Set the reference location name, latitude, longitude, and allowed radius (metres). "
        "A 'Use My Current Location' button auto-fills coordinates from the browser's GPS. "
        "The current configuration is displayed as JSON at the bottom."
    ),
]

for title_, desc in pages_desc:
    heading(title_, 2)
    para(desc)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. KEY CODE SNIPPETS
# ══════════════════════════════════════════════════════════════════════════════
heading("7. Key Code Snippets")

heading("7.1  Face Recognition — Recognize Faces in Frame", 2)
para("The core recognition function detects all faces in a captured image and matches them against stored encodings:")
code_block("""\
def recognize_faces(image, tolerance=0.5):
    img_array = np.array(image.convert("RGB"))
    locations = face_recognition.face_locations(img_array, model="hog")
    encodings = face_recognition.face_encodings(img_array, locations)

    known = load_encodings()       # {user_id: [enc, ...]}
    results = []

    for enc, loc in zip(encodings, locations):
        best_id, best_conf = None, 0
        for uid, stored_encs in known.items():
            distances = face_recognition.face_distance(stored_encs, enc)
            min_dist  = float(np.min(distances))
            if min_dist <= tolerance:
                conf = round((1 - min_dist) * 100, 1)
                if conf > best_conf:
                    best_id, best_conf = uid, conf
        results.append({"user_id": best_id, "location": loc,
                         "confidence": best_conf})
    return results, None""")

doc.add_paragraph()
heading("7.2  Face Encoding — Encode a Single Image", 2)
para("Used during registration to encode each captured training photo:")
code_block("""\
def encode_image(image):
    img_array = np.array(image.convert("RGB"))
    locations = face_recognition.face_locations(img_array)
    if len(locations) == 0:
        return None, "No face detected. Ensure you are well-lit and centred."
    if len(locations) > 1:
        return None, "Multiple faces detected. Only one person should be in frame."
    encoding = face_recognition.face_encodings(img_array, locations)[0]
    return encoding, None""")

doc.add_paragraph()
heading("7.3  Haversine Distance — GPS Location Check", 2)
para("Calculates the straight-line distance between two GPS coordinates for location-based attendance restriction:")
code_block("""\
def haversine_distance(lat1, lon1, lat2, lon2) -> float:
    R = 6_371_000  # Earth radius in metres
    phi1, phi2 = math.radians(lat1), math.radians(lat2)
    dphi    = math.radians(lat2 - lat1)
    dlambda = math.radians(lon2 - lon1)
    a = (math.sin(dphi / 2) ** 2
         + math.cos(phi1) * math.cos(phi2) * math.sin(dlambda / 2) ** 2)
    return 2 * R * math.asin(math.sqrt(a))""")

doc.add_paragraph()
heading("7.4  Mark Attendance — Check-in Logic", 2)
para("Inserts or validates an attendance record with duplicate-prevention:")
code_block("""\
def mark_checkin(user_id, location="Lecture Hall"):
    conn = get_connection()
    today = date.today().isoformat()

    existing = conn.execute(
        "SELECT id, check_out_time FROM attendance "
        "WHERE user_id=? AND date=? AND location=?",
        (user_id, today, location),
    ).fetchone()

    if existing:
        if existing["check_out_time"] is None:
            conn.close()
            return "already_in", existing["id"]
        conn.close()
        return "already_done", existing["id"]

    cursor = conn.execute(
        "INSERT INTO attendance (user_id, date, check_in_time, location, status) "
        "VALUES (?, ?, ?, ?, 'Present')",
        (user_id, today, datetime.now().isoformat(), location),
    )
    conn.commit()
    new_id = cursor.lastrowid
    conn.close()
    return "success", new_id""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. LOCATION RESTRICTION FEATURE
# ══════════════════════════════════════════════════════════════════════════════
heading("8. Location Restriction Feature")
para(
    "This feature ensures attendance can only be marked when the device is physically present "
    "at or near the designated classroom or campus location."
)
doc.add_paragraph()

heading("How It Works", 2)
steps = [
    "Admin opens Settings (page 6) and sets the reference GPS coordinates and allowed radius (default 100 m).",
    "Coordinates are saved as a singleton row in the location_settings table.",
    "When a user opens Mark Attendance, a Location Verification section appears.",
    "The user clicks 'Get Location' — the browser requests GPS permission and returns coordinates.",
    "The Haversine formula calculates the distance (metres) between the device and the reference point.",
    "If distance ≤ radius → green confirmation, attendance proceeds normally.",
    "If distance > radius → red error, the attendance button is blocked.",
    "If no reference location is configured → location check is skipped (backwards compatible).",
]
for i, s in enumerate(steps, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(s)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. SCREENSHOTS
# ══════════════════════════════════════════════════════════════════════════════
heading("9. Application Screenshots")

BASE = os.path.dirname(os.path.abspath(__file__))
SS = os.path.join(BASE, "screenshots")
ROOT = BASE

screenshot_sets = [
    # (path, caption)
    (os.path.join(SS, "1_login_page.png"),        "Figure 1 — Login Page (Student & Teacher tabs)"),
    (os.path.join(SS, "2_student_tab.png"),        "Figure 2 — Student Login Tab"),
    (os.path.join(SS, "3_wrong_pass.png"),         "Figure 3 — Invalid Credentials Error"),
    (os.path.join(SS, "register_tab.png"),         "Figure 4 — Registration Tab (new account)"),
    (os.path.join(SS, "register_form.png"),        "Figure 5 — Student Registration Form with Camera"),
    (os.path.join(SS, "4_student_dashboard.png"),  "Figure 6 — Student Dashboard (personal attendance view)"),
    (os.path.join(SS, "s1_home.png"),              "Figure 7 — Teacher / Admin Dashboard (class overview)"),
    (os.path.join(SS, "s3_mark.png"),              "Figure 8 — Mark Attendance Page (camera + status panel)"),
    (os.path.join(SS, "s4_blocked.png"),           "Figure 9 — Attendance Blocked (outside allowed radius)"),
    (os.path.join(SS, "s2_reports.png"),           "Figure 10 — Attendance Reports with Charts"),
    (os.path.join(SS, "student_attendance.png"),   "Figure 11 — Student Attendance Records View"),
    # Fallback to root-level screenshots if screenshots/ ones are missing
    (os.path.join(ROOT, "ss_manage.png"),          "Figure 12 — Manage Users Page"),
    (os.path.join(ROOT, "screenshot_dashboard.png"), "Figure 13 — Teacher Dashboard (alternate view)"),
    (os.path.join(ROOT, "ss_register2.png"),       "Figure 14 — Register Page with Enrolled Students Grid"),
    (os.path.join(ROOT, "ss_reg_db.png"),          "Figure 15 — Student Database Grid View"),
]

added = 0
for path, caption in screenshot_sets:
    if os.path.exists(path):
        heading(caption, 2)
        add_screenshot(path, caption)
        doc.add_paragraph()
        added += 1

if added == 0:
    para("No screenshot files were found. Place PNG/JPG files in the screenshots/ folder and re-run.")

# ── Save ──────────────────────────────────────────────────────────────────────
out = os.path.join(BASE, "Face_Recognition_Attendance_Documentation.docx")
doc.save(out)
print(f"Saved: {out}")
